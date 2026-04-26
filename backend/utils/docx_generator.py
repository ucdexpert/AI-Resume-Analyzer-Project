from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_resume_docx(data):
    doc = Document()
    
    # 1. Header
    name = doc.add_heading(data.get('full_name', 'Your Name'), 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = f"{data.get('location', '')} | {data.get('phone', '')} | {data.get('email', '')}"
    contact.add_run(contact_info).font.size = Pt(10)
    
    if data.get('linkedin'):
        links = doc.add_paragraph()
        links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links.add_run(f"LinkedIn: {data.get('linkedin')}").font.size = Pt(9)

    # 2. Summary
    if data.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data.get('summary')).style = 'Body Text'

    # 3. Experience
    if data.get('experience'):
        doc.add_heading('Work Experience', level=1)
        for exp in data.get('experience', []):
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('company', '')}").bold = True
            p.add_run(f"\t{exp.get('dates', '')}").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p2 = doc.add_paragraph()
            p2.add_run(f"{exp.get('job_title', '')}").italic = True
            
            # Bullet points
            desc = exp.get('description', '')
            if desc:
                bullets = desc.split('\n')
                for bullet in bullets:
                    if bullet.strip():
                        # Remove leading dots if any
                        clean_bullet = bullet.strip().lstrip('•').strip()
                        doc.add_paragraph(clean_bullet, style='List Bullet')

    # 4. Education
    if data.get('education'):
        doc.add_heading('Education', level=1)
        for edu in data.get('education', []):
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('school', '')}").bold = True
            p.add_run(f"\t{edu.get('dates', '')}").italic = True
            doc.add_paragraph(edu.get('degree', ''))

    # 5. Skills
    if data.get('skills'):
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(', '.join(data.get('skills', [])))

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()
